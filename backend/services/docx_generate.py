from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from core.config import report_root
from schemas.analyzer import GlobalAnalysisResponse
from schemas.common import MessageAnalysisReport


def set_font(run, font_name: str = "等线"):
    """统一设置 run 的中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc, text: str, level: int):
    """
    添加标题并统一设置中文字体：
    - 一级标题：18pt，居中
    - 二级标题：14pt
    """
    heading = doc.add_heading(text, level=level)

    for run in heading.runs:
        set_font(run)

    if level == 1:
        for run in heading.runs:
            run.font.size = Pt(18)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        for run in heading.runs:
            run.font.size = Pt(14)

    return heading


def generate_docx_report(
    name: str,
    created_at: datetime,
    situation: str,
    global_analysis: GlobalAnalysisResponse,
    message_analyses: List[MessageAnalysisReport],
) -> Path:
    """
    生成口语对话分析报告 DOCX 文件
    """
    doc = Document()

    # 设置默认段落样式为新宋体，字号12pt
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "新宋体"
    style_normal._element.rPr.rFonts.set(qn("w:eastAsia"), "新宋体")
    style_normal.font.size = Pt(12)

    # 添加标题
    add_heading(doc, "口语对话分析报告", level=1)

    # 添加基本信息
    doc.add_paragraph(f"时间: {created_at.strftime('%Y/%m/%d %H:%M:%S')}")
    doc.add_paragraph(f"情景：{situation}")

    # 总体分析
    add_heading(doc, "总体分析", level=2)
    doc.add_paragraph(f"语法、词汇: {global_analysis.grammar_analysis}")
    doc.add_paragraph(f"发音: {global_analysis.pronunciation_analysis}")
    doc.add_paragraph(f"表达: {global_analysis.expression_analysis}")

    # 对话内容与分析
    add_heading(doc, "对话内容与分析", level=2)

    for message in message_analyses:
        paragraph = doc.add_paragraph()
        run_index = paragraph.add_run(f"({message.index}) ")
        set_font(run_index)
        run_index.bold = True

        role_text = "你：" if message.role == "user" else "AI："
        run_content = paragraph.add_run(role_text + message.content)
        set_font(run_content)

        if message.role == "user":
            doc.add_paragraph(f"语法、词汇分析： {message.grammar_analysis}")
            doc.add_paragraph(f"发音分析： {message.pronunciation_analysis}")

    # 保存文件
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    docx_name = f"{name}.docx"
    final_path = report_root / docx_name
    with open(final_path, "wb") as f:
        f.write(output.read())

    return final_path


if __name__ == "__main__":
    name_1 = "test"
    situation_1 = "自由对话"
    created_at_1 = datetime.now()
    global_analysis_1 = GlobalAnalysisResponse(
        grammar_analysis="语法分析结果",
        pronunciation_analysis="发音分析结果",
        expression_analysis="表达分析结果",
    )
    message_analyses_1 = [
        MessageAnalysisReport(
            index=1,
            role="user",
            content="你好！",
            grammar_analysis="语法分析结果",
            pronunciation_analysis="发音分析结果",
        ),
        MessageAnalysisReport(
            index=2,
            role="assistant",
            content="你好！我能帮你什么？",
            grammar_analysis=None,
            pronunciation_analysis=None,
        ),
    ]
    path = generate_docx_report(
        name=name_1,
        created_at=created_at_1,
        situation=situation_1,
        global_analysis=global_analysis_1,
        message_analyses=message_analyses_1,
    )
